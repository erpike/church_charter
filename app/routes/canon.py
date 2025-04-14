from io import BytesIO

from docx import Document
from flask import Blueprint, abort, render_template, send_file

from storages.database import db
from storages.models import Canon, CanonChapter, CanonItem

# Create Blueprint
canon_bp = Blueprint("canon", __name__, url_prefix="/canon")


@canon_bp.route("/<int:canon_id>")
def detail(canon_id):
    """Display details of a specific canon."""
    with db:
        canon = Canon.get_or_none(Canon.id == canon_id)
        if canon is None:
            abort(404)

        # Get sorted content using the get_data method
        content = canon.get_data()

        return render_template(
            "canon_detail.html",
            canon=canon,
            content=content,
        )


@canon_bp.route("/<int:canon_id>/download-docx")
def download_docx(canon_id):
    """Download canon as a .docx file."""
    with db:
        canon = Canon.get_or_none(Canon.id == canon_id)
        if canon is None:
            abort(404)

        # Create a new Document
        doc = Document()

        # Add title
        doc.add_heading(canon.name, 0)

        # Add chapters and their items
        for chapter in canon.chapters.order_by(CanonChapter.position):
            # Add chapter title
            doc.add_heading(chapter.title, level=1)

            # Add items
            for item in chapter.items.order_by(CanonItem.position):
                if item.type == "hirmos":
                    doc.add_paragraph(f"Ирмос: {item.text}")
                elif item.type == "refrain":
                    doc.add_paragraph(f"Припев: {item.text}")
                else:
                    doc.add_paragraph(item.text)

            # Add a separator between chapters
            doc.add_paragraph()

        # Save the document to a BytesIO object
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        # Return the document as a downloadable file
        return send_file(
            doc_io,
            mimetype=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
            as_attachment=True,
            download_name=f"{canon.name}.docx",
        )
